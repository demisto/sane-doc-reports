from typing import Tuple, Union

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def _insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


class CellObject(object):
    """ An object containning a cell and it's inner:
     - paragraph (holds: runs (w:p element))
     - run (holds: text, pictures, text-styling (font))
     """

    def __init__(self, cell, add_run=True):
        self.cell = cell

        cell_paragraph, cell_run = self._get_cell_wrappers(add_run=add_run)
        self.paragraph = cell_paragraph
        self.run = cell_run

    def _get_cell_wrappers(self, add_run=True) -> Tuple[Paragraph, Union[Run, None]]:
        """
        Return the cell's paragraph and create a run object too, return them
        both. They are used to inject elements into the table cell.
        Run object:
        - https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects
        Paragraph Object:
        - https://python-docx.readthedocs.io/en/latest/api/text.html#paragraph-objects
        """
        paragraphs = self.cell.paragraphs
        paragraph = paragraphs[0]
        run = None
        if add_run:
            run = paragraph.add_run()
        return paragraph, run

    def add_run(self) -> None:
        self.run = self.paragraph.add_run()

    def get_last_paragraph(self) -> Paragraph:
        return self.cell.paragraphs[-1]

    def add_paragraph(self, style=None, add_run=True) -> Paragraph:
        self.paragraph = _insert_paragraph_after(self.paragraph, style=style)
        if add_run:
            self.run = self.paragraph.add_run()
        return self.paragraph
