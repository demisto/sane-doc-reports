from math import floor

from docx.shared import Pt

from sane_doc_reports.domain.CellObject import CellObject
from sane_doc_reports.domain.Element import Element
from sane_doc_reports.conf import DEBUG, TREND_MAIN_NUMBER_FONT_SIZE, \
    TREND_SECOND_NUMBER_FONT_SIZE, ALIGN_RIGHT, DEFAULT_DURATION_TITLE, \
    STYLE_KEY, PYDOCX_FONT_SIZE, DEFAULT_DURATION_TITLE_FONT_SIZE, \
    DEFAULT_DURATION_FONT_SIZE, PYDOCX_FONT_BOLD, \
    DEFAULT_DURATION_LABEL_FONT_SIZE, DURATION_MINUTES_LABEL, \
    DURATION_HOURS_LABEL, DURATION_DAYS_LABEL
from sane_doc_reports.elements import error
from sane_doc_reports.populate.utils import insert_text


def format_number(num):
    return ('0' + str(num))[-2:]


class DurationElement(Element):

    def insert(self):
        if DEBUG:
            print("Adding duration...")

        contents = self.section.contents

        days = '0'
        hours = '0'
        minutes = '0'
        if contents:
            result = 0
            if len(contents) > 0 and isinstance(contents[0]['data'],
                                                list) and len(
                    contents[0]['data']) > 0:
                result = contents[0]['data'][0]

            days = floor(result / (3600 * 24))
            result -= days * 3600 * 24

            hours = floor(result / 3600)
            result -= hours * 3600

            minutes = floor(result / 60)

            days = format_number(days)
            hours = format_number(hours)
            minutes = format_number(minutes)

        # Split the table as so:
        # +-----------+
        # | Title     |
        # +-----------+
        # | H | M | S |
        # +---+---+---+
        # .cell(row, col)
        table = self.cell_object.cell.add_table(rows=2, cols=3)
        if DEBUG:
            table.style = 'Table Grid'

        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 2))

        title = DEFAULT_DURATION_TITLE
        if len(contents) > 0 and 'name' in contents[0] and contents[0][
            'name'] != '':
            title = contents['data']['name']

        title_style = {PYDOCX_FONT_SIZE: DEFAULT_DURATION_TITLE_FONT_SIZE}
        insert_text(title_cell, title, title_style)

        style = {
            PYDOCX_FONT_SIZE: DEFAULT_DURATION_FONT_SIZE,
            PYDOCX_FONT_BOLD: True
        }
        style_label = {
            PYDOCX_FONT_SIZE: DEFAULT_DURATION_LABEL_FONT_SIZE
        }

        days_cell = table.cell(1, 0)
        insert_text(days_cell, days, style)
        insert_text(days_cell, DURATION_DAYS_LABEL, style_label,
                    add_run=True)

        hours_cell = table.cell(1, 1)
        insert_text(hours_cell, hours, style)
        insert_text(hours_cell, DURATION_HOURS_LABEL, style_label,
                    add_run=True)

        minutes_cell = table.cell(1, 2)
        insert_text(minutes_cell, minutes, style)
        insert_text(minutes_cell, DURATION_MINUTES_LABEL, style_label,
                    add_run=True)


def invoke(cell_object, section):
    if section.type != 'duration':
        section.contents = f'Called duration but not duration -  [{section}]'
        return error.invoke(cell_object, section)

    DurationElement(cell_object, section).insert()
