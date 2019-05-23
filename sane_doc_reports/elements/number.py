from docx.shared import Pt

from sane_doc_reports.domain.CellObject import CellObject
from sane_doc_reports.domain.Element import Element
from sane_doc_reports.conf import DEBUG, TREND_MAIN_NUMBER_FONT_SIZE, \
    TREND_SECOND_NUMBER_FONT_SIZE, ALIGN_RIGHT
from sane_doc_reports.elements import error


class NumberElement(Element):

    def insert(self):
        if DEBUG:
            print("Adding number: ", self.section.contents)

        table = self.cell_object.cell.add_table(rows=1, cols=1)

        if DEBUG:
            table.style = 'Table Grid'

        # Add the main number
        inner_cell = table.cell(0, 0)

        main_number = CellObject(inner_cell)
        main_number.run.text = str(self.section.contents)
        main_number.run.font.size = Pt(TREND_MAIN_NUMBER_FONT_SIZE)
        main_number.run.font.bold = True
        main_number.paragraph.alignment = ALIGN_RIGHT

        # Add the title
        title_paragraph = inner_cell.add_paragraph()
        title_run = title_paragraph.add_run()
        title_run.text = str(self.section.extra['title'])
        title_run.font.size = Pt(TREND_SECOND_NUMBER_FONT_SIZE)
        title_run.font.bold = False
        title_paragraph.alignment = ALIGN_RIGHT


def invoke(cell_object, section):
    if section.type != 'number':
        section.contents = f'Called number but not number -  [{section}]'
        return error.invoke(cell_object, section)

    NumberElement(cell_object, section).insert()
