from docx.table import _Cell

from sane_doc_reports.domain.CellObject import CellObject
from sane_doc_reports.domain.Section import Section
from sane_doc_reports.elements import text


def insert_text_into_cell(cell: _Cell, text_value: str, style={},
                          add_run=False):
    cell_object = CellObject(cell, add_run=add_run)
    section = Section('text', text_value, style, {})
    text.invoke(cell_object, section)
