import base64
import re
import tempfile
from io import BytesIO
import importlib
from pathlib import Path
from typing import Tuple

from docx.oxml import OxmlElement
import matplotlib
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph
from matplotlib import pyplot as plt
from matplotlib import colors as mcolors

from sane_doc_reports import CellObject, Section
from sane_doc_reports.conf import STYLE_KEY, SIZE_H_INCHES, SIZE_W_INCHES, \
    DPI, DEFAULT_DPI, PYDOCX_FONT_SIZE, PYDOCX_FONT_NAME, \
    PYDOCX_FONT_BOLD, PYDOCX_FONT_STRIKE, PYDOCX_FONT_UNDERLINE, \
    PYDOCX_FONT_ITALIC, PYDOCX_FONT_COLOR, PYDOCX_TEXT_ALIGN

colors = dict(mcolors.BASE_COLORS, **mcolors.CSS4_COLORS)
DEFAULT_BAR_COLOR = '#999999'
CHART_COLORS = ['#E57373', '#FF1D1E', '#FF5000', '#E55100', '#D74315',
                '#F06292', '#FF3F81', '#F50057', '#C2195B', '#E91D63',
                '#AD1457', '#CE93D8', '#EA80FC', '#FA99D0', '#FD5BDE',
                '#D500F9', '#AA00FF', '#BA68C8', '#B287FE', '#9575CD',
                '#AB47BC', '#8E24AA', '#8052F3', '#9FA8DA', '#7C71F5',
                '#536DFE', '#5C6BC0', '#3F51B5', '#6200EA', '#A3C9FF',
                '#64B5F6', '#03B0FF', '#2196F3', '#2979FF', '#295AFF',
                '#B7E7FF', '#81D4FA', '#80DEEA', '#00B8D4', '#039BE5',
                '#0277BD', '#1AEADE', '#18DEE5', '#00E5FF', '#4DD0E1',
                '#4DB6AC', '#0097A7', '#0097A7', '#64DC17', '#00E676',
                '#00C853', '#20B358', '#4CAF50', '#C5FE01', '#ADE901',
                '#50EB07', '#AED581', '#8BC34A', '#69A636', '#EDFE41',
                '#FFEA00', '#FFD740', '#F9A825', '#FB8C00', '#FF7500',
                '#DBDBDB', '#CFD8DC', '#9EB6C3', '#B2B7B9', '#989898',
                '#90A4AE']


def name_to_rgb(color_name: str):
    hex_color = name_to_hex(color_name)
    return hex_to_rgb(hex_color)


def name_to_hex(color_name: str):
    """ Get the hex representation of a color name (CSS4) """
    return colors[color_name].lower()


def hex_to_rgb(hex_color: str):
    """
    Convert a hex color string into an RGBColor object (used in python-docx)
    """
    hex_color = hex_color.lstrip('#')
    return RGBColor(*[int(hex_color[i:i + 2], 16) for i in (0, 2, 4)])


def open_b64_image(image_base64):
    """
    Open a virtual image file from base64 format of image.
    """
    prefix_regex = r'^data:.*?;base64,'
    raw_base64 = re.sub(prefix_regex, '', image_base64)
    f = BytesIO()
    f.write(base64.b64decode(raw_base64))
    f.seek(0)
    return f


def insert_by_type(type: str, cell_object: CellObject,
                   section: Section):
    """ Call a docx elemnt's insert method """
    func = importlib.import_module(f'sane_doc_reports.docx.{type}')

    func.invoke(cell_object, section)


def _insert_paragraph_after(paragraph):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    return new_para


def add_run(cell_object):
    """ Insert a paragraph so we could add a new element"""
    cell_object.paragraph = _insert_paragraph_after(cell_object.paragraph)
    cell_object.run = cell_object.paragraph.add_run()
    return cell_object


def has_run(cell_object: CellObject):
    """ A helper used to make sure to add a run """
    if cell_object.run is None:
        cell_object.add_run()


def plot(func):
    """ A decorator used to clear and resize each chart """

    def wrapper(*args, **kwargs):
        plt.clf()
        func(*args, **kwargs)

    return wrapper


def plt_t0_b64(plt: matplotlib.pyplot):
    """ Matplotlib to base64 url """
    path = Path(tempfile.mkdtemp()) / Path(
        next(tempfile._get_candidate_names()) + '.png')

    plt.savefig(str(path), format='png', bbox_inches='tight', figsize=(1, 1),
                dpi=80)

    with open(str(path), "rb") as f:
        img_base64 = base64.b64encode(f.read()).decode("utf-8", "ignore")
        b64 = f'data:image/png;base64,{img_base64}'

    path.unlink()
    return b64


def convert_plt_size(section: Section):
    """ Convert the plot size from pixels to word """
    size_w, size_h, dpi = (SIZE_W_INCHES, SIZE_H_INCHES, DPI)
    if 'dimensions' in section.layout:
        h = section.layout['dimensions']['height'] / DEFAULT_DPI
        w = section.layout['dimensions']['width'] / DEFAULT_DPI
        size_w, size_h, dpi = (w, h, DEFAULT_DPI)

    return size_w, size_h, dpi


def get_saturated_colors():
    """ Return named colors that are clearly visible on a white background """
    return [name for name, _ in colors.items()
            if 'light' not in name and 'white' not in name]


def apply_styling(cell_object, style):
    apply_cell_styling(cell_object, style)
    apply_paragraph_styling(cell_object, style)


def apply_cell_styling(cell_object, style):
    # Font size
    if PYDOCX_FONT_SIZE in style:
        cell_object.run.font.size = Pt(style[PYDOCX_FONT_SIZE])

    # Font family
    if PYDOCX_FONT_NAME in style:
        cell_object.run.font.name = style[PYDOCX_FONT_NAME]

    # Other characteristics
    if PYDOCX_FONT_BOLD in style:
        cell_object.run.font.bold = style[PYDOCX_FONT_BOLD]
    if PYDOCX_FONT_STRIKE in style:
        cell_object.run.font.strike = style[PYDOCX_FONT_STRIKE]
    if PYDOCX_FONT_UNDERLINE in style:
        cell_object.run.font.underline = style[PYDOCX_FONT_UNDERLINE]
    if PYDOCX_FONT_ITALIC in style:
        cell_object.run.font.italic = style[PYDOCX_FONT_ITALIC]

    # Font color
    if PYDOCX_FONT_COLOR in style:
        if style[PYDOCX_FONT_COLOR][0] != '#':
            cell_object.run.font.color.rgb = name_to_rgb(
                style[PYDOCX_FONT_COLOR])
        else:
            cell_object.run.font.color.rgb = hex_to_rgb(
                style[PYDOCX_FONT_COLOR])


def apply_paragraph_styling(cell_object, style):
    if PYDOCX_TEXT_ALIGN in style:
        if style[PYDOCX_TEXT_ALIGN] == 'left':
            cell_object.paragraph.paragraph_format.alignment = 0
        elif style[PYDOCX_TEXT_ALIGN] == 'right':
            cell_object.paragraph.paragraph_format.alignment = 2
        elif style[PYDOCX_TEXT_ALIGN] == 'center':
            cell_object.paragraph.paragraph_format.alignment = 1


def _hash_simple_value(s):
    """ djb2 """
    hash = 5381
    i = len(s)
    for _ in s:
        i = i - 1
        hash = (hash * 33) ^ ord(s[i])
    return hash & 0xFFFFFFFF


def get_chart_color(value):
    """ Trying to copy the sane-report color scheme """
    if not value:
        return DEFAULT_BAR_COLOR

    index = _hash_simple_value(value) % len(CHART_COLORS)
    return CHART_COLORS[index]


def get_ax_location(legend_style):
    align = legend_style.get('align', None)
    vertical_align = legend_style.get('verticalAlign', None)

    if not align or not vertical_align:
        return 'best'

    vertical_align = vertical_align.replace('top', 'upper').replace(
        'bottom', 'lower')
    return f'{vertical_align} {align}'


def get_colors(section_layout, objects):
    """ Return the chart colors and replace the default colors if they
    are hardcoded """
    default_colors = [get_chart_color(i) for i in objects]
    if not "legend" in section_layout or not isinstance(
            section_layout['legend'], list):
        return default_colors

    legend_colors = section_layout['legend']
    defined_colors = [i['name'] for i in legend_colors]
    ret_colors = []
    for name in objects:
        if name in defined_colors:
            ret_colors.append(
                legend_colors[defined_colors.index(name)]['color'])
        else:
            ret_colors.append(default_colors.pop())
    return ret_colors


def get_current_li(extra, list_type) -> Tuple[str, int, str]:
    """ Return the current list item style and indent level """
    list_level = 1
    list_type = list_type if 'list_type' not in extra else extra['list_type']
    p_style = list_type
    if 'list_level' in extra:
        list_level = int(extra['list_level']) + 1

        # Word doesn't have more than 3 levels of indentation
        if list_level >= 4:
            list_level = 3
        p_style = f'{list_type} {list_level}'
    return p_style, list_level, list_type


def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
